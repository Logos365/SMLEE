"""발표자료_오타점검_결과.md → Word 문서 변환"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ─── 색상 팔레트 ───────────────────────────────────────────────
RED       = RGBColor(0xC0, 0x00, 0x00)
DARKRED   = RGBColor(0x80, 0x00, 0x00)
GREEN     = RGBColor(0x37, 0x56, 0x23)
NAVY      = RGBColor(0x1F, 0x4E, 0x79)
BLUE      = RGBColor(0x2E, 0x75, 0xB6)
ORANGE    = RGBColor(0xED, 0x7D, 0x31)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)
DARKGRAY  = RGBColor(0x40, 0x40, 0x40)
LIGHTBLUE = RGBColor(0xBD, 0xD7, 0xEE)
LIGHTRED  = RGBColor(0xFF, 0xCC, 0xCC)
LIGHTYELLOW=RGBColor(0xFF, 0xF2, 0xCC)
LIGHTGREEN= RGBColor(0xE2, 0xEF, 0xDA)
LIGHTGRAY = RGBColor(0xF2, 0xF2, 0xF2)

def rgb_hex(rgb: RGBColor) -> str:
    """RGBColor → 6자리 헥스 문자열"""
    return f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'

def set_cell_bg(cell, rgb: RGBColor):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val','single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color','000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    if level == 1:
        p.paragraph_format.left_indent = Cm(0)
    else:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(14 if level==1 else 12 if level==2 else 11)
    # 하단 테두리 (h1만)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), rgb_hex(NAVY))
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_body(doc, text, indent=0.5, size=10.5, color=DARKGRAY, bold=False):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    return p

def add_bullet(doc, text, indent=1.0, size=10.5, marker='•', color=DARKGRAY):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f'{marker}  {text}')
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_badge_para(doc, badge, text, badge_color=RED, text_color=DARKGRAY, size=10.5):
    """🔴 난제... 스타일 배지 단락"""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f' {badge} ')
    r1.bold = True
    r1.font.size = Pt(size)
    r1.font.color.rgb = WHITE
    # 배지 배경 (highlight 없이 shading 이용 — run shading)
    rPr = r1._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    hc = rgb_hex(badge_color)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hc)
    rPr.append(shd)
    r2 = p.add_run(f'  {text}')
    r2.bold = True
    r2.font.size = Pt(size)
    r2.font.color.rgb = badge_color
    return p

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_table_header(table, headers, bg=NAVY, fg=WHITE, sizes=None):
    hdr = table.rows[0]
    for i, (cell, txt) in enumerate(zip(hdr.cells, headers)):
        set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        run.bold = True
        run.font.color.rgb = fg
        run.font.size = Pt(sizes[i] if sizes else 10)

# ═══════════════════════════════════════════════════════════════
doc = Document()

# ─── 페이지 여백 설정 ─────────────────────────────────────────
sec = doc.sections[0]
sec.page_width  = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.0)

# ─── 기본 스타일 ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ═══════════════════════════════════════════════════════════════
# 표지 헤더
# ═══════════════════════════════════════════════════════════════
tbl_cover = doc.add_table(rows=1, cols=1)
tbl_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = tbl_cover.rows[0].cells[0]
set_cell_bg(cell, NAVY)
cell.width = Cm(16)
p_title = cell.paragraphs[0]
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(10)
p_title.paragraph_format.space_after  = Pt(4)
r = p_title.add_run('발표자료 오타 및 용어 불일치 점검 결과')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = WHITE
r.font.name = '맑은 고딕'
r.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p_sub = cell.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(2)
p_sub.paragraph_format.space_after  = Pt(10)
for txt, sz, clr in [
    ('점검 대상: ', 10, LIGHTBLUE),
    ('20260609_첨단소재_원천기술 개발_Merge-V05.pdf  (50페이지)', 10, WHITE),
]:
    r2 = p_sub.add_run(txt)
    r2.font.size = Pt(sz); r2.font.color.rgb = clr
    r2.font.name = '맑은 고딕'
    r2.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p_info = cell.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_info.paragraph_format.space_after = Pt(10)
for txt, clr in [('점검일: 2026-06-11   |   ', LIGHTBLUE),
                  ('ALD 파트(기계연·대성기계) 전용 점검 포함', WHITE)]:
    r3 = p_info.add_run(txt)
    r3.font.size = Pt(9.5); r3.font.color.rgb = clr
    r3.font.name = '맑은 고딕'
    r3.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_paragraph()  # 여백

# ─── 범례 ─────────────────────────────────────────────────────
p_legend = doc.add_paragraph()
p_legend.paragraph_format.space_before = Pt(4)
p_legend.paragraph_format.space_after  = Pt(8)
p_legend.paragraph_format.left_indent  = Cm(0.5)
for badge, txt, clr in [
    ('🔴', '즉시 수정 필요  ', RED),
    ('🟡', '권장 수정  ', ORANGE),
    ('🟢', '확인 권장  ', GREEN),
]:
    r = p_legend.add_run(f'{badge} {txt}   ')
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = clr
    r.font.name = '맑은 고딕'
    r.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 1. 명백한 오타 — 즉시 수정
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '1. 명백한 오타 (즉시 수정 필요)', level=1)

errors = [
    {
        'no': '오타 1',
        'badge_color': RED,
        'title': '"hybide" → "hybrid"',
        'location': '14페이지 (03| 연구개발목표)',
        'before': 'O3-P2hybide활물질용량180mAh/g',
        'after':  'O3-P2 hybrid 활물질',
        'note':   '영문 오타. 심사위원에게 부정적 인상을 줄 수 있음',
    },
    {
        'no': '오타 2',
        'badge_color': RED,
        'title': '"나트튬" → "나트륨"',
        'location': '16페이지 (03| 연구개발목표 및 내용)',
        'before': 'O3-P2기반나트튬공침공정최적화',
        'after':  'O3-P2 기반 나트륨 공침 공정 최적화',
        'note':   '핵심 용어 오타. 제목급 텍스트이므로 즉시 수정 필요',
    },
    {
        'no': '오타 3',
        'badge_color': RED,
        'title': '"레서피" → "레시피" (2회 발생)',
        'location': '19페이지 (03| 연구개발목표 및 내용)',
        'before': '① 코팅양극재전기화학성능검증레서피이전및연속식공\n② ALD 공정레서피개발(3-4종화합물)',
        'after':  '레시피 이전 및 연속식 공정 / ALD 공정 레시피 개발',
        'note':   '동일 페이지에 2회 반복 발생',
    },
    {
        'no': '오타 4',
        'badge_color': RED,
        'title': '"직관적조장" → "직관적 조작"',
        'location': '38페이지 (별첨 05| 기관별 핵심보유기술 — 씨오알엔)',
        'before': '직관적조장, 사용편의성상승',
        'after':  '직관적 조작, 사용 편의성 향상',
        'note':   '"조장"은 "조작"의 오타로 추정',
    },
    {
        'no': '오타 5',
        'badge_color': RED,
        'title': '"Mxene" → "MXene"',
        'location': '44페이지 (별첨 05| 기관별 핵심보유기술 — 기계연)',
        'before': 'TiS₂/TiO₂@Mxene (Adv. Funct. Mater., 2024)',
        'after':  'TiS₂/TiO₂@MXene',
        'note':   'MXene은 고유 브랜드명으로 대소문자 구분 필수 (학계 표준: MXene)',
    },
    {
        'no': '오타 6',
        'badge_color': RED,
        'title': '"소듐이온전지" → "나트륨이온전지" (용어 불일치)',
        'location': '20페이지, 29페이지 (씨오알엔 관련 섹션)',
        'before': '소듐이온전지 (20p, 29p에 반복)',
        'after':  '나트륨이온전지 (전체 통일)',
        'note':   '김현수 박사님 p.12 코멘트와 직결. 동일 자료 내 혼용은 준비 부족으로 인식됨',
    },
]

for err in errors:
    add_badge_para(doc, err['no'], err['title'],
                   badge_color=err['badge_color'])

    # 위치
    p_loc = doc.add_paragraph()
    p_loc.paragraph_format.left_indent = Cm(1.2)
    p_loc.paragraph_format.space_before = Pt(1)
    p_loc.paragraph_format.space_after  = Pt(1)
    r = p_loc.add_run('위치: ')
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY
    r.font.name = '맑은 고딕'; r.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    r2 = p_loc.add_run(err['location'])
    r2.font.size = Pt(10); r2.font.color.rgb = DARKGRAY
    r2.font.name = '맑은 고딕'; r2.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 수정 전/후 테이블
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.rows[0].cells[0].width = Cm(2.5)
    tbl.rows[0].cells[1].width = Cm(13.0)

    labels = ['수정 전', '수정 후']
    contents = [err['before'], err['after']]
    bg_colors = [LIGHTRED, LIGHTGREEN]
    lbl_colors = [RED, GREEN]

    for ri, (lbl, cont, bg, lc) in enumerate(zip(labels, contents, bg_colors, lbl_colors)):
        c0 = tbl.rows[ri].cells[0]
        c1 = tbl.rows[ri].cells[1]
        set_cell_bg(c0, bg)
        set_cell_bg(c1, LIGHTGRAY if ri == 0 else LIGHTGREEN)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = c0.paragraphs[0].add_run(lbl)
        r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = lc
        r0.font.name = '맑은 고딕'; r0.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        r1 = c1.paragraphs[0].add_run(cont)
        r1.font.size = Pt(10); r1.font.color.rgb = DARKGRAY
        r1.font.name = '맑은 고딕'; r1.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 비고
    p_note = doc.add_paragraph()
    p_note.paragraph_format.left_indent  = Cm(1.2)
    p_note.paragraph_format.space_before = Pt(2)
    p_note.paragraph_format.space_after  = Pt(8)
    r_n = p_note.add_run('비고: ')
    r_n.bold = True; r_n.font.size = Pt(9.5); r_n.font.color.rgb = NAVY
    r_n.font.name = '맑은 고딕'; r_n.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    r_n2 = p_note.add_run(err['note'])
    r_n2.font.size = Pt(9.5); r_n2.font.color.rgb = DARKGRAY
    r_n2.font.name = '맑은 고딕'; r_n2.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ═══════════════════════════════════════════════════════════════
# 2. 권장 수정 사항
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading(doc, '2. 권장 수정 사항 (용어·표현 불일치)', level=1)

recommended = [
    {
        'no': '표현 불일치 1',
        'badge_color': ORANGE,
        'title': '"고출력" → "고전압" (RFP 키워드 부합)',
        'location': '26페이지 (05| 사업화전략 및 기대효과 — 대성기계)',
        'before': '고출력·장수명나트륨이온전지구현',
        'after':  '고전압(4.4V급)·장수명 나트륨이온전지 구현',
        'note':   '김현수 박사님 코멘트 "rfp에 있는 고전압 특성 강조" 명시. RFP 핵심 키워드 직접 반영 필요',
    },
    {
        'no': '표현 불일치 2',
        'badge_color': ORANGE,
        'title': '기술이전 금액 단위: "건" → "억원"',
        'location': '15페이지 (03| 연구개발목표)',
        'before': '기술이전 금액 단위: 건',
        'after':  '기술이전 금액 단위: 억원',
        'note':   '김현수 박사님 코멘트 "기술이전 금액의 단위는 건에서 억원으로 수정" — 반드시 수정',
    },
    {
        'no': '표현 불일치 3',
        'badge_color': ORANGE,
        'title': 'V₂O₃ → V₂O₅ 확인 필요',
        'location': '18페이지',
        'before': 'V2O3, TiO2, NbO5의 Precursor...',
        'after':  'V₂O₅ (오산화바나듐), Nb₂O₅ (오산화이나이오브) 확인 후 수정',
        'note':   '본 과제 대상 소재는 V₂O₅. V₂O₃(삼산화바나듐)는 다른 소재. 화학식 오류는 심사위원 신뢰도 훼손',
    },
]

for rec in recommended:
    add_badge_para(doc, rec['no'], rec['title'],
                   badge_color=rec['badge_color'])
    p_loc = doc.add_paragraph()
    p_loc.paragraph_format.left_indent = Cm(1.2)
    p_loc.paragraph_format.space_before = Pt(1)
    p_loc.paragraph_format.space_after  = Pt(1)
    r = p_loc.add_run('위치: ')
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY
    r.font.name = '맑은 고딕'; r.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    r2 = p_loc.add_run(rec['location'])
    r2.font.size = Pt(10); r2.font.color.rgb = DARKGRAY
    r2.font.name = '맑은 고딕'; r2.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    tbl2 = doc.add_table(rows=2, cols=2)
    tbl2.style = 'Table Grid'
    for ri, (lbl, cont, bg) in enumerate(
        [('수정 전', rec['before'], LIGHTRED),
         ('수정 후', rec['after'],  LIGHTGREEN)]):
        c0 = tbl2.rows[ri].cells[0]
        c1 = tbl2.rows[ri].cells[1]
        set_cell_bg(c0, bg)
        set_cell_bg(c1, LIGHTGRAY if ri==0 else LIGHTGREEN)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        lc = RED if ri==0 else GREEN
        r0 = c0.paragraphs[0].add_run(lbl)
        r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = lc
        r0.font.name = '맑은 고딕'; r0.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        r1 = c1.paragraphs[0].add_run(cont)
        r1.font.size = Pt(10); r1.font.color.rgb = DARKGRAY
        r1.font.name = '맑은 고딕'; r1.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    p_note = doc.add_paragraph()
    p_note.paragraph_format.left_indent = Cm(1.2)
    p_note.paragraph_format.space_before = Pt(2)
    p_note.paragraph_format.space_after  = Pt(8)
    r_n = p_note.add_run('비고: ')
    r_n.bold = True; r_n.font.size = Pt(9.5); r_n.font.color.rgb = NAVY
    r_n.font.name = '맑은 고딕'; r_n.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    r_n2 = p_note.add_run(rec['note'])
    r_n2.font.size = Pt(9.5); r_n2.font.color.rgb = DARKGRAY
    r_n2.font.name = '맑은 고딕'; r_n2.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ═══════════════════════════════════════════════════════════════
# 3. 확인 권장 사항
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading(doc, '3. 확인 권장 사항', level=1)

checks = [
    ('확인 1', '연구목표 최종목표 문구 (셀 용량 0.8→1.0 Ah, 용량유지율 동일 여부)',
     '15페이지',
     '김현수 박사님 코멘트 "0.8Ah에서 1.0Ah 증가가 의미 있는 목표치인지?" — 발표 시 근거 설명 준비 필요'),
    ('확인 2', '씨오알엔 "셀 조립라인 투자 예상 60억원", "단계적 확장" 문구 수정',
     '27~29페이지',
     '김현수 박사님 "현재 기구축되어 있다는 전제하에 과제 진행이 바람직함" — 문구 삭제 또는 긍정적으로 전환'),
    ('확인 3', '씨오알엔 발표자료 페이지 수 1페이지로 통일',
     '27페이지',
     '김현수 박사님 "다른 기관과 동일하게 1페이지로 통일" — 분량 조정 필요'),
    ('확인 4', '별첨 슬라이드 "0X" 번호 실제 번호로 수정',
     '45~50페이지 (별첨)',
     '"0X" 임시 번호 그대로 제출 시 완성도 낮아 보임. 실제 번호로 교체 필요'),
]
for no, title, loc, note in checks:
    add_badge_para(doc, no, title, badge_color=GREEN)
    p_l = doc.add_paragraph()
    p_l.paragraph_format.left_indent = Cm(1.2)
    p_l.paragraph_format.space_before = Pt(1)
    p_l.paragraph_format.space_after  = Pt(1)
    r = p_l.add_run('위치: '); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=NAVY
    r.font.name='맑은 고딕'; r.font.element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')
    r2 = p_l.add_run(loc); r2.font.size=Pt(10); r2.font.color.rgb=DARKGRAY
    r2.font.name='맑은 고딕'; r2.font.element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')
    p_n = doc.add_paragraph()
    p_n.paragraph_format.left_indent=Cm(1.2)
    p_n.paragraph_format.space_before=Pt(2); p_n.paragraph_format.space_after=Pt(8)
    rn=p_n.add_run('비고: '); rn.bold=True; rn.font.size=Pt(9.5); rn.font.color.rgb=NAVY
    rn.font.name='맑은 고딕'; rn.font.element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')
    rn2=p_n.add_run(note); rn2.font.size=Pt(9.5); rn2.font.color.rgb=DARKGRAY
    rn2.font.name='맑은 고딕'; rn2.font.element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')

# ═══════════════════════════════════════════════════════════════
# 4. ALD 파트 점검 결과
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading(doc, '4. ALD 파트 (기계연·대성기계) 전용 점검 결과', level=1)

p_ok = doc.add_paragraph()
p_ok.paragraph_format.left_indent = Cm(0.5)
p_ok.paragraph_format.space_before = Pt(4)
p_ok.paragraph_format.space_after  = Pt(6)
r_ok = p_ok.add_run('ALD 파트 자체 오타 없음.  ')
r_ok.bold = True; r_ok.font.size = Pt(11); r_ok.font.color.rgb = GREEN
r_ok.font.name = '맑은 고딕'; r_ok.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
r_ok2 = p_ok.add_run('오타는 타 기관 파트(씨오알엔, 전북대 등) 섹션에 주로 집중되어 있음.')
r_ok2.font.size = Pt(10.5); r_ok2.font.color.rgb = DARKGRAY
r_ok2.font.name = '맑은 고딕'; r_ok2.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

ald_pages = [
    ('slide1 (p.3)', 'ALD 코팅 소재 역할 비교, 문헌 인용', '정상'),
    ('slide2 (p.5)', '기관 역할 분담 설명', '정상'),
    ('slide3 (p.6)', '간트차트, 마일스톤', '정상'),
    ('slide4 (p.8)', 'RTD 알고리즘 표기 (μ-2σ ≥ t_pulse)', '정상'),
    ('slide5 (p.9)', 'KPI 단위 (Ω·cm²)', '정상'),
    ('slide6 (p.12)', '레시피 표기', '정상'),
    ('slide7 (p.15)', '논문 타임라인', '정상'),
    ('slide8 (p.24)', '기술 강점 요약', '정상'),
]
tbl_ald = doc.add_table(rows=len(ald_pages)+1, cols=3)
tbl_ald.style = 'Table Grid'
add_table_header(tbl_ald, ['슬라이드', '점검 항목', '결과'], bg=NAVY, fg=WHITE, sizes=[10,10,10])
for ri, (pg, item, result) in enumerate(ald_pages, 1):
    row = tbl_ald.rows[ri]
    bg = LIGHTGREEN if result == '정상' else LIGHTRED
    for ci, txt in enumerate([pg, item, result]):
        set_cell_bg(row.cells[ci], bg if ci == 2 else LIGHTGRAY)
        p = row.cells[ci].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(txt)
        r.font.size = Pt(10)
        r.font.color.rgb = GREEN if (ci==2 and result=='정상') else DARKGRAY
        r.bold = (ci == 2)
        r.font.name = '맑은 고딕'; r.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ═══════════════════════════════════════════════════════════════
# 5. 수정 우선순위 요약
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading(doc, '5. 수정 우선순위 요약', level=1)

priority_data = [
    ('1', '"나트튬" → "나트륨"', '16p', '핵심 용어 오타', RED),
    ('2', '"레서피" → "레시피"', '19p (×2)', '반복 오타', RED),
    ('3', '"hybide" → "hybrid"', '14p', '영문 오타', RED),
    ('4', '"소듐이온전지" → "나트륨이온전지"', '20p, 29p', '용어 불일치', RED),
    ('5', '"고출력" → "고전압"', '26p', 'RFP 부합성', ORANGE),
    ('6', '"Mxene" → "MXene"', '44p', '학술 표기 오류', RED),
    ('7', '"직관적조장" → "직관적 조작"', '38p', '의미 오류', RED),
    ('8', '기술이전 금액 단위 (건→억원)', '15p', '심사 기준 부적합', ORANGE),
    ('9', 'V₂O₃ → V₂O₅ 확인', '18p', '화학식 오류 가능', ORANGE),
    ('10', '별첨 "0X" 번호 수정', '45~50p', '완성도', GREEN),
]

tbl_pri = doc.add_table(rows=len(priority_data)+1, cols=5)
tbl_pri.style = 'Table Grid'
add_table_header(tbl_pri,
    ['순위', '항목', '위치', '영향도', '중요도'],
    bg=NAVY, fg=WHITE, sizes=[9,10,9,10,9])

col_widths = [1.5, 7.0, 2.0, 4.0, 2.0]
for ci, w in enumerate(col_widths):
    for row in tbl_pri.rows:
        row.cells[ci].width = Cm(w)

for ri, (rank, item, loc, impact, clr) in enumerate(priority_data, 1):
    row = tbl_pri.rows[ri]
    bg_row = LIGHTRED if clr==RED else LIGHTYELLOW if clr==ORANGE else LIGHTGREEN
    for ci, (txt, align) in enumerate([
        (rank,   WD_ALIGN_PARAGRAPH.CENTER),
        (item,   WD_ALIGN_PARAGRAPH.LEFT),
        (loc,    WD_ALIGN_PARAGRAPH.CENTER),
        (impact, WD_ALIGN_PARAGRAPH.LEFT),
        ('즉시 수정' if clr==RED else '권장' if clr==ORANGE else '확인',
                 WD_ALIGN_PARAGRAPH.CENTER),
    ]):
        c = row.cells[ci]
        set_cell_bg(c, bg_row if ci in [0,4] else LIGHTGRAY)
        p = c.paragraphs[0]
        p.alignment = align
        r = p.add_run(txt)
        r.font.size = Pt(9.5)
        r.font.color.rgb = clr if ci in [0,4] else DARKGRAY
        r.bold = ci in [0, 4]
        r.font.name = '맑은 고딕'; r.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ─── 하단 주석 ────────────────────────────────────────────────
doc.add_paragraph()
p_foot = doc.add_paragraph()
p_foot.paragraph_format.left_indent = Cm(0.5)
p_foot.paragraph_format.space_before = Pt(12)
r_f = p_foot.add_run(
    '※ 본 점검은 PDF 텍스트 추출 기반으로 수행되었습니다. '
    'PDF에서 이미지로 삽입된 텍스트(그림 속 글자 등)는 점검에서 제외될 수 있으니 육안 확인을 병행하시기 바랍니다.')
r_f.font.size = Pt(9); r_f.font.color.rgb = DARKGRAY; r_f.italic = True
r_f.font.name = '맑은 고딕'; r_f.font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ─── 저장 ─────────────────────────────────────────────────────
output = (r"d:\23_GitHub\01_Office_PC\030_Project\032_Proposed Project"
          r"\08_2026년 첨단소재 사업화 기술 개발(성과확장형)\발표자료_오타점검_결과.docx")
doc.save(output)
print("Saved:", output)
