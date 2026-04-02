from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 페이지 여백 설정 (A4)
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── 헬퍼 함수
def set_font(run, name='맑은 고딕', size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.insert(0, rFonts)

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sizes  = {1: 14, 2: 12, 3: 11}
    colors = {1: (0x1F, 0x49, 0x7D), 2: (0x26, 0x63, 0x99)}
    set_font(run, size=sizes.get(level, 11), bold=True, color=colors.get(level))
    return p

def add_body(doc, text, indent=0, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=10, bold=bold)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=9, color=(0x44, 0x44, 0x44))
    return p

def add_figure_box(doc, fig_num):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'D9E1F2')
    pPr.append(shd)
    run = p.add_run(f'[ 그림 {fig_num} 삽입 위치 ]')
    set_font(run, size=11, bold=True, color=(0x1F, 0x49, 0x7D))
    return p

def make_table(doc, header_row, data_rows, col_widths=None):
    tbl = doc.add_table(rows=len(data_rows)+1, cols=len(header_row))
    tbl.style = 'Table Grid'
    # 헤더
    for j, hdr in enumerate(header_row):
        cell = tbl.rows[0].cells[j]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(hdr)
        set_font(run, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, '1F497D')
    # 데이터
    for i, row_data in enumerate(data_rows):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i+1].cells[j]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            is_key = (j == 0)
            set_font(run, size=10, bold=is_key)
            if is_key:
                set_cell_shading(cell, 'D9E1F2')
    return tbl

# ══════════════════════════════════════════
#  제 목
# ══════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(4)
r = p_title.add_run('2026년도 하반기 미래개척융합과학기술개발사업\n기술수요조사서')
set_font(r, size=16, bold=True, color=(0x1F, 0x49, 0x7D))

p_sub = doc.add_paragraph()
p_sub.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(2)
r2 = p_sub.add_run('유망 신기술형 #4 – 구조형 배터리 복합소재')
set_font(r2, size=13, bold=True, color=(0x26, 0x63, 0x99))

p_info = doc.add_paragraph()
p_info.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
p_info.paragraph_format.space_after = Pt(10)
r3 = p_info.add_run('이승모 (한국기계연구원)  |  2026년 4월')
set_font(r3, size=10, color=(0x55, 0x55, 0x55))

# ══════════════════════════════════════════
#  1. 연구과제 개요
# ══════════════════════════════════════════
add_heading(doc, '1. 연구과제 개요', level=1)
make_table(doc,
    ['항목', '내용'],
    [
        ('과제명',     '금속 침투 공정 기반 다기능 구조형 배터리 복합소재 기술 개발'),
        ('기술 분야',  '신소재·나노 / 에너지저장 / 구조복합소재'),
        ('핵심 키워드','구조형 배터리, 탄소 직물 전극, 금속 침투(Metal Infiltration), MXene 복합소재, 다기능 복합재료'),
        ('연구 유형',  '도전·혁신형 기초원천 연구'),
    ]
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ══════════════════════════════════════════
#  2. 기술개발 필요성
# ══════════════════════════════════════════
add_heading(doc, '2. 기술개발 필요성', level=1)
add_heading(doc, '2.1 사회·산업적 배경', level=2)
add_body(doc,
    '전기차(EV), 도심항공모빌리티(UAM), 무인기(드론), 우주발사체 등 차세대 모빌리티 산업의 핵심 과제는 '
    '에너지 밀도 극대화와 시스템 경량화의 동시 달성이다. 현재 전기차 배터리 팩은 차량 전체 중량의 '
    '25~30%를 차지하며, 별도 구조 부재(차체 프레임, 언더바디)가 추가되어 이중 무게 부담이 발생한다. '
    '구조형 배터리 복합소재(Structural Battery Composite, SBC)는 탄소섬유 강화 복합재료의 역학적 '
    '골격 기능과 리튬이온 배터리의 에너지 저장 기능을 단일 소재 시스템 안에서 구현하여 시스템 중량을 '
    '30~40% 절감할 수 있다. 스웨덴 Chalmers 대학, MIT, DLR 등 세계 선도 기관들이 2020년대 이후 '
    '집중 투자하고 있으나, 기계적 물성과 전기화학 성능 사이의 상충 관계(trade-off) 극복이 핵심 난제이다.')

add_figure_box(doc, 1)
add_caption(doc,
    '그림 1. 구조형 배터리 복합소재의 멀티펑셔널 개념도\n'
    '(좌) 기존 전기차: 배터리 팩 + 구조 프레임 분리 구조 vs. '
    '(우) SBC: 하중 지지와 Li+ 이온 이동을 단일 소재에서 동시 구현 (중량 절감 30~40%)')

add_heading(doc, '2.2 기술적 배경 및 도전 과제', level=2)
add_body(doc, '구조형 배터리 복합소재 구현을 위한 세 가지 핵심 기술이 요구된다.')
for b in [
    '구조 전극 소재: 탄소섬유/직물이 집전체이자 하중 전달 매개체로 기능하며, 역학적 취약점 없이 고용량 활물질이 균일 탑재되어야 한다.',
    '구조용 전해질(Structural Electrolyte): 기계적 강성을 부여하는 고분자 매트릭스(에폭시, PMMA)가 동시에 Li+ 이온 전도체로 기능해야 한다.',
    '계면 설계: 탄소섬유-활물질-전해질 삼중 계면에서 기계적 접착력과 이온 전달 경로를 동시에 최적화해야 한다.',
]:
    add_body(doc, '  •  ' + b, indent=0.3)
add_body(doc,
    '현재 ALD 기반 금속 침투(Metal Infiltration) 공정을 활용한 계면 강화 접근법은 거의 시도되지 않은 '
    '혁신적 전략으로, 본 연구팀의 원천 기술과 직접 연계된다.')

# ══════════════════════════════════════════
#  3. 연구개발 목표 및 내용
# ══════════════════════════════════════════
add_heading(doc, '3. 연구개발 목표 및 내용', level=1)
add_heading(doc, '3.1 최종 목표', level=2)
add_body(doc,
    'ALD/CVD 기반 금속 침투 공정을 구조형 배터리 복합소재에 적용하여, 기계적 물성(탄성계수, 인장강도, '
    '층간전단강도)과 전기화학 성능(에너지 밀도, 사이클 수명)을 동시에 향상시키는 원천 소재 기술을 확보한다.')

make_table(doc,
    ['성능 지표', '목표값'],
    [
        ('구조 전극 탄성계수',  '>= 50 GPa  (탄소섬유 기준 유지율 90% 이상)'),
        ('면적당 에너지 밀도',  '>= 5 mAh/cm2'),
        ('사이클 수명',         '500회 후 용량 유지율 >= 80%'),
        ('층간전단강도 (ILSS)', '>= 30 MPa'),
    ]
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_heading(doc, '3.2 연구 내용', level=2)
add_heading(doc, '1단계 (1~2년차): 구조 전극 소재 원천 기술 개발', level=3)
for s in [
    'ALD 기반 금속산화물(TiO2, Al2O3)/황화물(TiS2) 나노층을 탄소 직물 섬유 표면에 원자 수준으로 균일 코팅',
    '금속 침투 공정을 통한 탄소 직물 내부 공극 선택적 강화: 섬유-매트릭스 계면의 기계적 맞물림(interlocking) 구조 형성',
    'MXene(Ti3C2Tx)과 전이금속 황화물의 하이브리드 구조를 탄소 직물 상에 통합하여 고용량 구조 전극 구현',
]:
    add_body(doc, '  •  ' + s, indent=0.3)

add_figure_box(doc, 2)
add_caption(doc,
    '그림 2. ALD/금속 침투 공정을 이용한 탄소 직물 기반 구조 전극 제조 공정 모식도\n'
    '(1) 탄소 직물 원본  →  (2) ALD 나노층 코팅  →  (3) 금속 침투(계면 맞물림)  →  (4) MXene/TiS2 활물질 탑재')

add_heading(doc, '2단계 (3년차): 구조용 전해질 및 계면 최적화', level=3)
for s in [
    '에폭시/이온성액체 복합 전해질 시스템 개발: Li+ 이온 전도도 >= 10-4 S/cm 달성',
    '탄소섬유-활물질-전해질 삼중 계면의 기계/전기화학적 특성 인시투(in-situ) 라만 분광 분석',
    '적층 구조 최적화를 통한 구조형 배터리 단위 셀 제작 및 성능 검증',
]:
    add_body(doc, '  •  ' + s, indent=0.3)

add_figure_box(doc, 3)
add_caption(doc,
    '그림 3. 구조형 배터리 복합소재의 층 구조 및 인시투 라만 분석 개념도\n'
    '(좌) 단위 셀 적층 단면: 양극 탄소직물 전극 / 구조용 전해질층 / 음극 탄소직물 전극\n'
    '(우) 인장 하중 + 전기화학 충방전 동시 인가 실험 셋업 및 라만 스펙트럼 변화')

# ══════════════════════════════════════════
#  4. 연구팀 역량
# ══════════════════════════════════════════
add_heading(doc, '4. 연구팀 역량 및 보유 기술', level=1)
add_heading(doc, '4.1 핵심 역량', level=2)

competencies = [
    ('1.  금속 침투 기반 복합소재 역학 강화 (세계 최초 원천기술)',
     'ALD를 이용한 거미줄/콜라겐 내부로의 금속 침투를 통해 기계적 물성을 획기적으로 향상시키는 기술을 세계 최초로 개발하였다 [1, 2]. '
     '단순 코팅이 아닌 소재 내부로의 금속 침투로 역학적 맞물림 구조를 형성하여 층간 접착력과 파괴 인성을 동시에 향상시키는 이 기술은 '
     '구조형 배터리의 탄소섬유 전극-매트릭스 계면 강화에 직접 적용 가능하다.'),
    ('2.  탄소 직물 기반 고성능 에너지 저장 소재',
     '면(cotton) 직물의 금속산화물 보조 탄화(carbonization)로 슈퍼커패시터 전극을 구현 [3]하였으며, '
     'VC/V2O3 복합 나노입자를 탑재한 플렉시블 리튬-황 배터리 전극 기술을 확보하였다 [4, 5].'),
    ('3.  MXene 기반 고성능 배터리 및 고전도성 복합소재',
     'TiS2/TiO2@MXene 하이브리드 호스트를 개발하여 리튬-황 배터리의 황 산화환원 반응을 가속화하였으며 [6], '
     'Pt 침투 MXene의 전기/열전도도를 획기적으로 향상시켰다 [7].'),
    ('4.  2D 나노소재 복합체 구조-물성 상관관계',
     '자유-standing 그래핀 산화물 페이퍼의 두께 의존적 역학 특성 규명 [8] 및 CVD 그래핀의 파괴 특성 연구 [9]를 통해 '
     '나노소재 복합체의 역학적 거동에 대한 깊은 이해를 보유하고 있다.'),
]
for title, body in competencies:
    add_body(doc, title, bold=True)
    add_body(doc, body, indent=0.5)

add_figure_box(doc, 4)
add_caption(doc,
    '그림 4. 연구팀의 구조형 배터리 복합소재 관련 핵심 연구 성과 맵\n'
    '[1] 거미줄 금속 침투 (Science 2009)  [2] 탄소 직물 Li-S 배터리 (Energy Storage Materials 2022)\n'
    '[3] TiS2/TiO2@MXene (Adv. Funct. Mater. 2023)  [4] Pt 침투 MXene (JMCA 2021)\n'
    '[5] 면 직물 에너지 텍스타일 (Carbon 2019)  [6] 인시투 라만 분석 (ACS AMI 2014)')

add_heading(doc, '4.2 연구팀 구성 및 인프라', level=2)
for item in [
    '책임연구원: 이승모 박사 (한국기계연구원) — SCI 논문 90편 이상, 피인용 3,000회 이상, h-index 28',
    '공동연구기관: 탄소섬유 복합재 구조 해석 전문 그룹 (항공우주 분야 협력)',
    '주요 보유 장비: ALD 시스템, 라만 분광기, 만능재료시험기(UTM), EIS/CV/GCD 전기화학 분석, SEM/TEM, XRD',
]:
    add_body(doc, '  •  ' + item, indent=0.3)

# ══════════════════════════════════════════
#  5. 기대효과
# ══════════════════════════════════════════
add_heading(doc, '5. 연구개발 기대효과', level=1)
effects = [
    ('기술적 기대효과', [
        '세계 최초 ALD/금속 침투 기반 구조형 배터리 복합소재 원천 기술 확보',
        '기계적 물성과 전기화학 성능 동시 최적화 설계 원리 규명',
        '탄소 직물 기반 구조 전극의 역할을 에너지 저장 구조 부재로 확장',
    ]),
    ('경제·산업적 기대효과', [
        '전기차 배터리 팩+구조 부재 통합으로 차량 중량 5~10% 절감 → 주행 거리 연장',
        'UAM, 드론 등 초경량 모빌리티 시스템의 에너지 자율성 향상',
        '구조형 배터리 소재 글로벌 시장(2030년 약 7억 달러 예상) 선점 기여',
    ]),
    ('사회적 기대효과', [
        '전기차 주행거리 불안 해소 및 탄소 중립 모빌리티 전환 가속화',
        '경량·고에너지 드론 기술을 통한 물류·재난 대응 능력 향상',
    ]),
]
for category, items in effects:
    add_body(doc, '▶ ' + category, bold=True)
    for item in items:
        add_body(doc, '  •  ' + item, indent=0.3)

# ══════════════════════════════════════════
#  6. 참고문헌
# ══════════════════════════════════════════
add_heading(doc, '6. 참고문헌', level=1)
refs = [
    '[1]  S.M. Lee et al., "Greatly increased toughness of infiltrated spider silk," Science 324(5926), 488-492, 2009. (피인용 503회)',
    '[2]  S.M. Lee et al., "Improved mechanical stability of dried collagen membrane after metal infiltration," ACS Appl. Mater. Interfaces 2(8), 2436-2441, 2010. (피인용 99회)',
    '[3]  S. Won, H.C. Shim, J.H. Kim, S.M. Lee, "Turning cotton into tough energy textile via metal oxide assisted carbonization," Carbon 153, 257-264, 2019. (피인용 18회)',
    '[4]  D.V. Lam et al., "Calligraphic ink enabling washable conductive textile electrodes for supercapacitors," J. Mater. Chem. A 4, 4082-4088, 2016. (피인용 40회)',
    '[5]  V.P. Nguyen et al., "Porous carbon textile decorated with VC/V2O3-x hybrid nanoparticles: Dual-functional host for flexible Li-S full batteries," Energy Storage Materials 46, 542-552, 2022. (피인용 65회)',
    '[6]  V.P. Nguyen et al., "Accelerated sulfur evolution reactions by TiS2/TiO2@MXene host for high-volumetric-energy-density lithium-sulfur batteries," Adv. Funct. Mater. 33(35), 2303503, 2023. (피인용 135회)',
    '[7]  V.P. Nguyen et al., "Drastically increased electrical and thermal conductivities of Pt-infiltrated MXenes," J. Mater. Chem. A 9(17), 10739-10746, 2021. (피인용 49회)',
    '[8]  T. Gong, D.V. Lam et al., "Thickness dependence of the mechanical properties of free-standing graphene oxide papers," Adv. Funct. Mater. 25, 3756-3763, 2015. (피인용 101회)',
    '[9]  Y. Hwangbo, C.K. Lee, S.M. Kim et al., "Fracture characteristics of monolayer CVD-graphene," Scientific Reports 4(1), 4439, 2014. (피인용 125회)',
    '[10] S.M. Lee et al., "In situ Raman spectroscopic study of Al-infiltrated spider dragline silk under tensile deformation," ACS Appl. Mater. Interfaces 6(19), 16827-16834, 2014. (피인용 43회)',
    '[11] V.P. Nguyen et al., "Intercalation-conversion hybrid cathode enabled by MXene-driven TiO2/TiS2 heterostructure for high-energy-density Li-S battery," Small Structures 5(8), 2400196, 2024. (피인용 26회)',
    '[12] D.V. Lam et al., "Carbon textile decorated with pseudocapacitive VC/VxOy for high-performance flexible supercapacitors," Small 13, 1702702, 2017. (피인용 32회)',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(ref)
    set_font(run, size=9)

# 저장
out_path = r'D:\23_GitHub\01_Office_PC\030_Project\032_Proposed Project\06_미래융합파이오니어_수요조사\수요조사서_구조형배터리복합소재.docx'
doc.save(out_path)
print('저장 완료:', out_path)
