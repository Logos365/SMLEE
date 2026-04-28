from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Default style
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)


def heading(doc, text, size=12, bold=True, before=8, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    return p


def para(doc, text, size=10, indent=0, before=2, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    return p


def bullet(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('• ' + text)
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    return p


def fig_caption(doc, label, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(label + ' ')
    r1.font.bold = True
    r1.font.size = Pt(9.5)
    r1.font.name = '맑은 고딕'
    r2 = p.add_run(desc)
    r2.font.size = Pt(9)
    r2.font.name = '맑은 고딕'
    r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ===== COVER =====
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[별첨1. 예비 연구계획서(연구내용) 파일 업로드 서식]')
run.font.size = Pt(10)
run.font.name = '맑은 고딕'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026년도 핵심연구(전략형) 신규과제')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = '맑은 고딕'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('예비 연구계획서(연구내용)')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = '맑은 고딕'

doc.add_paragraph()

tbl = doc.add_table(rows=4, cols=2)
tbl.style = 'Table Grid'
tbl.columns[0].width = Cm(3.0)
tbl.columns[1].width = Cm(13.5)

data = [
    ('유형', '혁신전략-04'),
    ('아젠다', '리튬이온전지 및 핵심소재'),
    ('국문 과제명', 'ALD 기반 불화물 인공 SEI 설계를 통한 리튬 금속 음극의 균일 핵생성 메커니즘 규명 및 덴드라이트 억제 원리 연구'),
    ('영문 과제명', 'Elucidation of Uniform Lithium Nucleation Mechanism and Dendrite Suppression Principles via ALD-based Artificial SEI Design for Li Metal Anodes'),
]
for i, (k, v) in enumerate(data):
    tbl.rows[i].cells[0].text = k
    tbl.rows[i].cells[1].text = v
    for cell in tbl.rows[i].cells:
        for pp in cell.paragraphs:
            for run in pp.runs:
                run.font.size = Pt(9.5)
                run.font.name = '맑은 고딕'

doc.add_page_break()

# ===== SECTION 1 =====
heading(doc, '1. 사업목적과 아젠다와의 부합성', size=12)

heading(doc, '1) 사업목적과의 부합성', size=11, bold=True, before=4, after=3)
para(doc,
     '리튬(Li) 금속 음극은 흑연 대비 약 10배의 이론 용량(3,860 mAh/g)을 가지며, 차세대 고에너지밀도 이차전지(350 Wh/kg급 이상)를 실현하는 핵심 소재이다. 그러나 충·방전 시 반복적으로 발생하는 덴드라이트(dendrite) 성장과 불균일한 SEI(solid electrolyte interphase) 형성은 안전성 저하·수명 열화의 근본 원인으로, 실용화의 가장 높은 기술적 장벽을 형성하고 있다.')
para(doc,
     '이 문제를 원자 수준에서 해결하는 계면 제어 연구는 나노 수준의 물질 합성·분석 역량과 정밀한 전기화학적 해석을 동시에 요구하는 심층 기초연구 영역이다. 기업 R&D에서 단기 상용화 관점에서 접근하기 어려운 이 연구는, 국가 차원의 장기적·체계적 기초연구 지원이 반드시 필요하다. 본 연구는 원자층증착(ALD) 기반 독보적 원천 기술을 바탕으로 Li 금속 음극 계면의 핵생성 메커니즘을 규명하는 기초연구로서, 핵심연구(전략형) 사업의 취지와 정확히 부합한다.')

heading(doc, '2) 아젠다와의 부합성', size=11, bold=True, before=4, after=3)
para(doc,
     '본 연구는 아젠다 「혁신전략-04: 리튬이온전지 및 핵심소재」에 해당한다. 아젠다 정의에 명시된 "현 리튬이온전지의 에너지밀도 극대화(350 Wh/kg급) 및 가격경쟁력·공정 친환경화 확보를 위한 리튬이온전지 소재 및 공정 기술"에 직접적으로 대응한다.')
para(doc,
     'ALD 공정은 ① 기상(vapor-phase) 반응에 의해 무용매·저온 증착이 가능하고, ② 1 Å 수준의 박막 두께 정밀 제어를 통해 불순물 없는 고품질 계면 형성이 가능하며, ③ 기존 습식 코팅 대비 공정 친환경성이 우수하다. 연구 대상인 ZnF2, AlF3 등 불화물계 ALD 박막은 Li+ 이온 전도성이 우수하고 전기화학적 환원 조건에서 LiF-rich SEI로 전환되는 특성을 지녀, 친환경적 배터리용 음극 소재 공정 기술 개발에 직접 기여한다.')

# ===== SECTION 2 =====
heading(doc, '2. 연구주제의 적합성', size=12)

heading(doc, '2-1) 연구 개념 및 핵심어', size=11, bold=True, before=4, after=3)
para(doc,
     '본 연구의 핵심 개념은 "ALD 공정으로 집전체(Cu foil) 또는 Li 금속 음극 표면에 불화물계 나노 박막을 원자 수준으로 정밀 코팅하여, Li+ 이온의 균일한 플럭스(flux)를 유도하고 덴드라이트-프리(dendrite-free) Li 증착을 실현하는 것"이다.')
para(doc,
     '핵심어: 원자층증착(ALD) | 인공 SEI(artificial SEI) | 불화물 코팅(ZnF2/AlF3/LiF) | Li 핵생성(nucleation) 제어 | 덴드라이트 억제 | Li 금속 음극 | 이차전지')

fig_caption(doc, '[그림 1]',
            '연구 개념도. (좌) ALD 코팅 없는 Cu 집전체에서의 불균일 Li 핵생성 및 덴드라이트 성장 문제 모식도; '
            '(중) ALD ZnF2 증착 공정 사이클(전구체 노출 → 퍼지 → 반응물 노출 → 퍼지) 및 박막 형성 메커니즘; '
            '(우) 불화물 ALD 인공 SEI 형성 후 균일 Li 핵생성 및 평탄한 Li 증착 결과. '
            '레이아웃: 가로 3분할 패널, 화살표로 문제→해결 흐름 표시, 각 패널 하단에 SEM 이미지 배치 공간 확보.')

heading(doc, '2-2) 연구 가설 및 도출 근거', size=11, bold=True, before=4, after=3)
para(doc,
     '【연구 가설】 "ALD 공정으로 증착된 나노 두께(5~100 nm)의 불화물계 박막은 집전체 표면 에너지를 균일화하여 Li 핵생성 과전압을 저감하고, 충·방전 중 전기화학적 활성화 과정에서 LiF-rich 인공 SEI로 변환되어 Li+ 이온 균일 플럭스를 유지함으로써 장기적인 덴드라이트 억제를 달성할 수 있다."',
     before=2, after=3)

bullet(doc, '연구책임자 그룹은 ALD ZnF2 박막이 Cu 집전체 위에서 Li 핵생성 과전압을 현저히 낮추고 균일한 Li 증착을 유도함을 실험적으로 확인하였다 (VP Nguyen et al., Adv. Sci. 2025).')
bullet(doc, '90 nm ZnF2 ALD 박막이 Zn 금속 음극 안정화에도 유효함을 검증하여, 불화물 ALD의 범용 음극 보호 특성을 확인하였다 (VP Nguyen et al., ACS Energy Lett. 2025).')
bullet(doc, '문헌상 LiF-rich SEI는 균일한 Li+ 이온 전도도와 기계적 강도를 제공하여 덴드라이트 성장을 근본적으로 억제하는 것으로 확인되어 있다 [8,9,10].')
bullet(doc, '연구책임자는 20년 이상 ALD 공정 연구 경력(Science 2009 ALD/VPI 논문 포함)을 보유하며, 이차전지 ALD 적용 SCI 논문 다수 발표.')

heading(doc, '2-3) 연구의 필요성', size=11, bold=True, before=4, after=3)
para(doc,
     'Li 금속 음극 기반 고에너지밀도 전지는 전기차·에너지저장장치(ESS) 등 탄소중립 실현의 핵심 기술이다. 12대 국가전략기술 중 "이차전지" 분야는 글로벌 기술 패권 경쟁의 최전선이며, 350 Wh/kg 이상의 에너지밀도 달성을 위해 Li 금속 음극으로의 전환이 불가피하다. 그러나 덴드라이트 성장 및 무한 부피 팽창 문제가 실용화를 가로막고 있다. ALD 기반 정밀 계면 제어는 이 문제를 원자 수준에서 접근할 수 있는 차별화된 방법론이며, 대한민국이 선도적 원천 기술을 확보할 수 있는 전략적 연구 방향이다.')

# ===== SECTION 3 =====
heading(doc, '3. 연구목표의 도전성 및 실현 가능성', size=12)

heading(doc, '1) 연구과제 목표 및 실현 가능성', size=11, bold=True, before=4, after=3)
para(doc,
     '【최종 연구 목표】 ALD 기반 불화물계 인공 SEI를 활용하여 Li 금속 음극에서 균일 Li 증착을 달성하는 원자 수준 메커니즘을 규명하고, 에너지밀도 350 Wh/kg급 Li 금속 전지의 기초 소재·공정 원천 기술을 확립한다.',
     before=2)

tbl2 = doc.add_table(rows=3, cols=2)
tbl2.style = 'Table Grid'
tbl2.columns[0].width = Cm(3.5)
tbl2.columns[1].width = Cm(12.5)
tbl2.rows[0].cells[0].text = '보유 역량'
tbl2.rows[0].cells[1].text = '세부 내용'
tbl2.rows[1].cells[0].text = 'ALD 장비 및 전문성'
tbl2.rows[1].cells[1].text = '4인치 ALD 장비 2대 보유 (금속산화물·금속불화물·귀금속 증착 경험). 이차전지 ALD 소재 연구 SCI 논문 다수. ZnF2·Al2O3·TiO2·Pt 등 이차전지 실적 보유.'
tbl2.rows[2].cells[0].text = '분석 인프라'
tbl2.rows[2].cells[1].text = 'XPS, Cryo-TEM, EIS, in-situ SEM, Operando XRD 접근 가능. DFT 계산 협력 연구진 확보.'
for row in tbl2.rows:
    for cell in row.cells:
        for pp in cell.paragraphs:
            for run in pp.runs:
                run.font.size = Pt(9.5)
                run.font.name = '맑은 고딕'

para(doc, '', before=2, after=2)

heading(doc, '2) 연구과제의 내용', size=11, bold=True, before=4, after=3)

fig_caption(doc, '[그림 2]',
            '연구 로드맵 및 Task 구성도. 상단에 전체 연구 흐름(Task 1→2→3) 화살표 다이어그램, 하단에 4개 연차별 주요 연구 활동 간트 차트(Gantt chart). '
            '각 Task 박스 내에 핵심 분석법 아이콘(TEM/XPS/EIS/DFT) 배치. 레이아웃: 세로형, A4 절반 크기.')

heading(doc, 'Task 1. ALD 불화물 박막 합성 및 계면 특성 분석 (1~2년차)', size=10.5, bold=True, before=4, after=2)
bullet(doc, 'ZnF2, AlF3, LiF ALD 전구체·반응물 최적화 및 증착 조건(온도, 사이클 수, 박막 두께 5~100 nm) 체계적 탐색')
bullet(doc, 'XRD, XPS, ToF-SIMS, TEM/EDS를 이용한 결정구조·화학 조성·계면 결합 상태 분석')
bullet(doc, 'EIS 및 Li+ 이온 전도도 측정으로 이온 수송 특성 평가; 핀홀-프리(pinhole-free) 박막 합성 조건 확립')

heading(doc, 'Task 2. In-situ/Operando 분석 기반 Li 핵생성 메커니즘 규명 (2~3년차)', size=10.5, bold=True, before=4, after=2)
bullet(doc, 'Cryo-TEM 및 AFM을 이용한 초기 Li 핵생성(nucleation) 나노구조의 실시간 관찰')
bullet(doc, 'Operando XRD 및 in-situ SEM을 통한 충·방전 중 Li 증착/용해 거동 및 SEI 진화 모니터링')
bullet(doc, 'DFT(밀도함수이론) 계산으로 불화물/Li 계면에서의 Li+ 이온 확산 에너지 배리어 이론적 규명')
bullet(doc, '불화물 박막 두께·조성에 따른 Li 핵생성 과전압 및 증착 균일도 상관관계 데이터베이스 구축')

heading(doc, 'Task 3. Full Cell 평가 및 인공 SEI 설계 원칙 도출 (3~4년차)', size=10.5, bold=True, before=4, after=2)
bullet(doc, 'ALD 코팅된 Li 금속 음극을 적용한 Li||NCM 풀 셀 제작 및 장기 사이클(>300회) 성능 평가')
bullet(doc, 'Post-mortem 분석(XPS, TEM)으로 사이클 후 SEI 조성 변화 추적 및 열화 메커니즘 규명')
bullet(doc, 'ZnF2 vs AlF3 vs LiF 비교 분석을 통한 범용 "ALD 기반 인공 SEI 설계 원칙" 제안')
bullet(doc, '목표 성능: 쿨롱 효율(CE) >99.5%, 300사이클 후 용량 유지율 >80%')

# Year table
heading(doc, '【연차별 목표 요약】', size=10.5, before=4, after=2)
tbl3 = doc.add_table(rows=5, cols=3)
tbl3.style = 'Table Grid'
tbl3.columns[0].width = Cm(1.8)
tbl3.columns[1].width = Cm(8.5)
tbl3.columns[2].width = Cm(5.7)
hdr = tbl3.rows[0].cells
hdr[0].text = '연차'
hdr[1].text = '주요 연구 내용'
hdr[2].text = '핵심 성과물'
rows_data = [
    ('1년차', 'ZnF2/AlF3 ALD 증착 조건 최적화, XPS/TEM 박막 분석', 'ALD 레시피 확립, 박막 특성 데이터'),
    ('2년차', 'In-situ AFM 핵생성 관찰, EIS 분석, 박막 두께-성능 상관관계 규명', 'SCI 논문 1편, 핵생성 과전압 데이터베이스'),
    ('3년차', 'Operando XRD/Cryo-TEM, DFT 계산, Full cell 프로토타입 제작', 'SCI 논문 1편, 메커니즘 규명'),
    ('4년차', 'Full cell 장기 평가, Post-mortem 분석, 설계 원칙 도출', 'SCI 논문 1편, 특허 출원, 설계 원칙 제안'),
]
for i, (yr, cont, out) in enumerate(rows_data):
    tbl3.rows[i+1].cells[0].text = yr
    tbl3.rows[i+1].cells[1].text = cont
    tbl3.rows[i+1].cells[2].text = out
for row in tbl3.rows:
    for cell in row.cells:
        for pp in cell.paragraphs:
            for run in pp.runs:
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'

para(doc, '', before=2, after=2)
heading(doc, '【연구의 도전성】', size=10.5, before=4, after=2)
para(doc,
     'Li 금속 음극의 덴드라이트 억제는 전 세계적으로 아직 완전히 해결되지 않은 난제이다. 본 연구는 ALD의 원자 수준 정밀 제어를 이용하여 인공 SEI의 이온 전도 메커니즘을 근본적으로 규명한다는 점에서 도전적이며, in-situ/operando 분석과 DFT 계산을 결합한 다학제적 접근이 기존 연구와 차별화되는 핵심 특징이다. 성공 시 대한민국 배터리 소재 분야의 글로벌 원천 기술 선점에 기여한다.',
     before=2)

# References
doc.add_page_break()
heading(doc, '참고문헌 (Reference)', size=12, before=4, after=4)
refs = [
    '[1] VP Nguyen, HC Shim, YW Byeon, JH Kim, SM Lee*, "Taming Lithium Nucleation and Growth on Cu Current Collector by Electrochemical Activation of ZnF2 Layer", Adv. Sci. 12, 2416426 (2025).',
    '[2] VP Nguyen, M Park, YW Byeon, et al., "Ultrathin Yet Effective: 90 nm ZnF2 Layer for Stabilizing Zinc-Metal Anodes", ACS Energy Lett. 10, 5503-5511 (2025).',
    '[3] VP Nguyen, D Kim, J Kang, et al., "Heterostructured Sn:SnO2 Nanodots for High-Performance Li-S Batteries with Kinetics-Enhanced Cathode and Dendrite-Free Anode", Adv. Funct. Mater. 35, 2507991 (2025).',
    '[4] VP Nguyen, JS Park, HC Shim, JM Yuk, JH Kim, D Kim, SM Lee*, "Accelerated Sulfur Evolution Reactions by TiS2/TiO2@MXene Host for High-Volumetric-Energy-Density Lithium-Sulfur Batteries", Adv. Funct. Mater. 33, 2303503 (2023).',
    '[5] VP Nguyen, IH Kim, HC Shim, et al., "Porous Carbon Textile Decorated with VC/V2O3-X Hybrid Nanoparticles: Dual-functional Host for Flexible Li-S Full Batteries", Energy Storage Mater. 46, 542 (2022).',
    '[6] SM Lee, E Pippel, U Gosele, et al., "Greatly Increased Toughness of Infiltrated Spider Silk (ALD/VPI)", Science 324, 488-492 (2009).',
    '[7] DV Lam, VH Nguyen, H Yoo, et al., "Dry Synthesis of Sulfur-Terminated MXene as Multifunctional Catalyst for Stable Lithium-Sulfur Batteries", Small 21, 2411668 (2025).',
    '[8] D Lin, Y Liu, Y Cui, "Reviving the Lithium Metal Anode for High-Energy Batteries", Nat. Nanotechnol. 12, 194-206 (2017).',
    '[9] Y Li, Y Li, A Pei, et al., "Atomic Structure of Sensitive Battery Materials and Interfaces Revealed by Cryo-Electron Microscopy", Science 358, 506-510 (2017).',
    '[10] J Zhao, L Liao, F Shi, et al., "Surface Fluorination of Reactive Battery Anode Materials for Enhanced Stability", JACS 139, 11550-11558 (2017).',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(ref)
    run.font.size = Pt(9)
    run.font.name = '맑은 고딕'

out_path = '연구계획서_초안_이승모_ALD_Li음극.docx'
doc.save(out_path)
print('Saved:', out_path)
